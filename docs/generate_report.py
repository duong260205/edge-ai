#!/usr/bin/env python3
"""
Tạo file Word (.docx) hướng dẫn xây dựng Edge AI Model cho dự án 
Hệ thống Voice + Face Authentication trên ESP32-S3
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Set default East Asian font
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Arial')

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0A, 0x22, 0x40)
    hs.font.name = 'Arial'
    if level == 1:
        hs.font.size = Pt(22)
    elif level == 2:
        hs.font.size = Pt(16)
    elif level == 3:
        hs.font.size = Pt(13)
    elif level == 4:
        hs.font.size = Pt(11)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_code_block(doc, code_text, indent=0.5):
    """Add a formatted code block to the document"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), 'F5F5FA')
    shading.set(qn('w:color'), 'auto')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_note(doc, text, style_name='info'):
    """Add a highlighted note box"""
    p = doc.add_paragraph()
    if style_name == 'info':
        run = p.add_run('💡 ')
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    elif style_name == 'warning':
        run = p.add_run('⚠️ ')
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    elif style_name == 'success':
        run = p.add_run('✅ ')
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    return p

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        if bold or header:
            run.bold = True
        if header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:fill'), 'E3F2FD')
            cell._element.get_or_add_tcPr().append(shading)
    return row

# ============================================================
# COVER PAGE
# ============================================================
# Add spacing
for _ in range(4):
    doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HƯỚNG DẪN XÂY DỰNG\nEDGE AI MODEL')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x22, 0x40)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Voice + Face Authentication trên ESP32-S3')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

doc.add_paragraph()

# Info block
info_items = [
    ('📌 Dự án:', 'Hệ thống khóa cửa thông minh đa nhân tố'),
    ('🎯 Công nghệ:', 'TensorFlow Lite Micro, ESP-DL, ESP32-S3'),
    ('🔬 ML Models:', 'CNN Voice Classifier (13.6K params) + MobileFaceNet'),
    ('⚙️ Framework:', 'ESP-IDF v5.x, FreeRTOS, TFLite Micro'),
    ('📅 Phiên bản:', '1.0 — 2026'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label} ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (Manual)
# ============================================================
doc.add_heading('Mục Lục', level=1)
toc_items = [
    '1. Tổng Quan Hệ Thống',
    '2. Kiến Trúc Edge AI trên ESP32-S3',
    '3. Voice Authentication Pipeline',
    '   3.1. Thu thập dữ liệu',
    '   3.2. Feature Engineering — MFCC',
    '   3.3. Kiến trúc mô hình CNN',
    '   3.4. Huấn luyện & Đánh giá',
    '   3.5. Lượng tử hóa INT8',
    '   3.6. Triển khai TFLite Micro',
    '4. Face Recognition Pipeline',
    '   4.1. Pre-trained Models (ESP-DL)',
    '   4.2. Face Detection & Recognition',
    '   4.3. Enrollment & Database',
    '5. Dual-Auth State Machine',
    '6. Tối ưu & Troubleshooting',
    '7. Kết Luận',
    'Phụ Lục: Mã nguồn tham khảo',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        p.runs[0].bold = True

doc.add_page_break()

# ============================================================
# 1. TỔNG QUAN HỆ THỐNG
# ============================================================
doc.add_heading('1. Tổng Quan Hệ Thống', level=1)

doc.add_paragraph(
    'Dự án xây dựng một hệ thống xác thực đa nhân tố (Multi-Factor Authentication) '
    'chạy hoàn toàn trên vi điều khiển ESP32-S3. Hệ thống kết hợp hai phương thức '
    'sinh trắc học: nhận diện khuôn mặt (Face Recognition) và xác thực giọng nói '
    '(Voice Authentication).'
)

doc.add_heading('1.1. Thông số phần cứng', level=2)

hw_table = doc.add_table(rows=1, cols=3)
hw_table.style = 'Light List Accent 1'
hw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(hw_table, ['Thiết bị', 'Model', 'Kết nối'], header=True)
add_table_row(hw_table, ['Vi điều khiển', 'ESP32-S3 (LX7 dual-core @240MHz, 8MB PSRAM, 16MB Flash)', 'SoC chính'])
add_table_row(hw_table, ['Camera', 'OV2640 (UXGA, JPEG, 8-bit parallel)', '8-bit DVP'])
add_table_row(hw_table, ['Microphone', 'INMP441 (I2S, 16kHz, 16-bit)', 'I2S (BCLK=40, WS=41, DATA=42)'])
add_table_row(hw_table, ['Servo', 'SG90 (PWM 50Hz)', 'GPIO 47 (LEDC)'])
add_table_row(hw_table, ['LED', 'WS2812 RGB', 'GPIO 48 (RMT)'])

doc.add_paragraph()

doc.add_heading('1.2. Kiến trúc phần mềm', level=2)

doc.add_paragraph(
    'Hệ thống được xây dựng trên nền tảng ESP-IDF v5.x với FreeRTOS làm hệ điều hành thời gian thực. '
    'Các tác vụ chính bao gồm:'
)
tasks = [
    ('State Machine Task (priority 6):', 'Quản lý trạng thái hệ thống, điều phối luồng xác thực'),
    ('Voice Auth Task (priority 5):', 'Thu âm I2S, VAD, chạy TFLite inference liên tục'),
    ('Face Recognition Task (priority 5):', 'Chụp ảnh, phát hiện và nhận diện khuôn mặt (chỉ chạy khi cần)'),
    ('HTTP Server Task (priority 5):', 'Phục vụ web UI, REST API, MJPEG stream'),
]
for name, desc in tasks:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ' ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 2. KIẾN TRÚC EDGE AI
# ============================================================
doc.add_heading('2. Kiến Trúc Edge AI trên ESP32-S3', level=1)

doc.add_paragraph(
    'Edge AI là việc chạy các mô hình Machine Learning trực tiếp trên thiết bị biên (edge device) '
    'thay vì gửi dữ liệu lên cloud. Điều này mang lại nhiều lợi ích:'
)

benefits = [
    ('Bảo mật:', 'Dữ liệu không rời khỏi thiết bị — không lo rò rỉ'),
    ('Tốc độ:', 'Inference trong ~50ms, không phụ thuộc mạng'),
    ('Chi phí:', 'Không tốn phí cloud/server'),
    ('Offline:', 'Hoạt động 100% không cần internet'),
]
for name, desc in benefits:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ' ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.1. TFLite Micro Framework', level=2)
doc.add_paragraph(
    'TensorFlow Lite Micro là phiên bản rút gọn của TensorFlow Lite dành cho vi điều khiển. '
    'Đặc điểm:'
)
micro_features = [
    'Dung lượng tối thiểu: ~20KB cho interpreter',
    'Hỗ trợ INT8 quantization (giảm 4x kích thước)',
    'Không cần OS, không cần heap động phức tạp',
    'Tensor Arena cấp phát trước (40KB PSRAM)',
]
for f in micro_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.2. ESP-DL Framework', level=2)
doc.add_paragraph(
    'ESP-DL (Espressif Deep Learning) là thư viện AI của Espressif, cung cấp các model nhận diện khuôn mặt được '
    'train sẵn và tối ưu cho ESP32-S3. ESP-DL tận dụng:'
)
espdl_features = [
    'Vector extension (PIE) trên ESP32-S3 cho tính toán vector hóa',
    'Tối ưu convolution cho ma trận nhỏ',
    'Các model .espdl đã được compile sẵn, lưu trong Flash',
]
for f in espdl_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. VOICE AUTHENTICATION PIPELINE
# ============================================================
doc.add_heading('3. Voice Authentication Pipeline', level=1)

doc.add_paragraph(
    'Đây là bài toán Binary Classification: phân biệt giọng chủ nhà (target) '
    'với mọi âm thanh khác (noise). Pipeline đầy đủ:'
)

add_code_block(doc, """\
Audio Raw (16-bit int, 16000 samples)
  → [I2S Capture] → [Gain x4 + Clipping] → [VAD]
  → [Hann Window] → [Real FFT (512-point)] → [Mel-scale binning]
  → [Log Energy] → MFCC Tensor (32, 13, 1)
  → [Conv2D → MaxPool → Conv2D → MaxPool → Flatten → Dense → Sigmoid]
  → Score [0, 1] → Threshold ≥ 0.7?""")

# 3.1 Data Collection
doc.add_heading('3.1. Thu thập dữ liệu', level=2)

doc.add_paragraph(
    'Dữ liệu được thu thập qua file collect_data.py sử dụng thư viện sounddevice. '
    'Mỗi mẫu âm thanh dài 1 giây, lấy mẫu ở 16kHz, 16-bit.'
)

doc.add_heading('Cấu trúc thư mục dataset', level=3)
add_code_block(doc, """\
dataset/
├── target/         ← Giọng chủ nhà (class 1)
│   ├── sample_0.wav
│   ├── sample_1.wav
│   └── ... (20-30 files)
└── noise/          ← Tiếng ồn/người khác (class 0)
    ├── sample_0.wav
    ├── sample_1.wav
    └── ... (20-50 files, có thể import ESC-50)""")

doc.add_paragraph(
    'Có thể import thêm dữ liệu từ dataset ESC-50 (~1000 files, 29 categories) '
    'qua file select_esc50.py để tăng độ đa dạng cho class noise.'
)

add_note(doc, 'Data càng đa dạng → Model chống nhiễu càng tốt!', 'info')

doc.add_heading('Tiền xử lý trên ESP32', level=3)
doc.add_paragraph(
    'Tín hiệu từ microphone INMP441 được xử lý qua các bước:'
)
add_code_block(doc, """\
// audio_capture.cpp — I2S capture + gain + clipping
for (size_t i = 0; i < valid_samples; i++) {
    // INMP441: 24-bit left-justified in 32-bit slot
    int32_t sample = chunk_buf[i] >> 14;  // >>14 = gain x4

    // Clipping: chống tràn số
    if (sample > 32767) sample = 32767;
    else if (sample < -32768) sample = -32768;

    buffer[samples_read++] = (int16_t)sample;
}""")

# 3.2 MFCC
doc.add_heading('3.2. Feature Engineering — MFCC', level=2)

doc.add_paragraph(
    'MFCC (Mel-Frequency Cepstral Coefficients) là kỹ thuật trích xuất đặc trưng '
    'từ tín hiệu âm thanh, mô phỏng cách tai người cảm nhận âm thanh. '
    'Thay vì dùng raw audio (16,000 samples), ta chuyển thành tensor (32, 13) '
    '— giảm ~38 lần kích thước!'
)

doc.add_heading('Tham số MFCC', level=3)

mfcc_table = doc.add_table(rows=1, cols=3)
mfcc_table.style = 'Light List Accent 1'
add_table_row(mfcc_table, ['Tham số', 'Giá trị', 'Giải thích'], header=True)
add_table_row(mfcc_table, ['frames', '32', 'Chia 1s thành 32 khung (31.25ms/khung)'])
add_table_row(mfcc_table, ['hop_length', '512', '= 32ms tại 16kHz, overlap 18.75ms'])
add_table_row(mfcc_table, ['fft_size', '512', 'Độ phân giải tần số: 31.25Hz/bin'])
add_table_row(mfcc_table, ['n_mfcc', '13', '13 hệ số Mel (cân bằng accuracy/size)'])

doc.add_paragraph()

doc.add_heading('Quy trình trích MFCC (đồng bộ Python — C++)', level=3)
doc.add_paragraph(
    'Điểm quan trọng nhất: Feature extraction phải HOÀN TOÀN GIỐNG NHAU giữa training (Python) '
    'và inference (C++). Nếu không, model sẽ cho kết quả sai trên thiết bị thật.'
)

doc.add_paragraph('Python (train_quantize_export.py):', style='List Bullet')
add_code_block(doc, """\
def extract_mfcc(file_path):
    audio, sr = librosa.load(file_path, sr=16000)
    audio = audio / max(abs(audio))   # normalize

    frames, hop, fft_size = 32, 512, 512
    window = 0.5 * (1 - cos(2π*n/512))  # Hann window

    for i in range(32):
        segment = audio[i*512 : i*512+512]
        windowed = segment * window
        fft = np.fft.rfft(windowed)

        for b in range(13):  # 13 Mel bands
            energy = sum(|fft[b*19:(b+1)*19]|^2)
            features[i,b] = log(energy/512^2 + 1e-6)
    return features  # (32, 13)""")

doc.add_paragraph('C++ (voice_auth.cpp):', style='List Bullet')
add_code_block(doc, """\
static void compute_mfcc_like_features(..., TfLiteTensor* input) {
    const int frames = 32, hop = 512;
    float scale = input->params.scale;
    int zp = input->params.zero_point;

    for (int i = 0; i < frames; i++) {
        // Window + FFT (ESP-DSP hardware accelerated)
        dsps_fft2r_fc32(fft_buf, 512);
        dsps_bit_rev_fc32(fft_buf, 512);

        for (int b = 0; b < 13; b++) {
            float energy = sum(|fft[bin_start:bin_end]|^2);
            float val = logf(energy / (512*512) + 1e-6f);
            // Quantize to INT8
            int8_t q = roundf(val / scale) + zp;
            input->data.int8[i*13 + b] = q;
        }
    }
}""")

add_note(doc, 'Lưu ý: Code dùng bin_start = b*(256/13) thay vì mel-scale triangular filters — đây là MFCC approximate, đơn giản hóa cho embedded.', 'warning')

# 3.3 Model Architecture
doc.add_heading('3.3. Kiến trúc mô hình CNN', level=2)

doc.add_paragraph(
    'Mô hình sử dụng CNN (Convolutional Neural Network) với kiến trúc nhẹ, '
    'chỉ 13,601 tham số (~12KB sau INT8 quantization).'
)

doc.add_heading('Chi tiết các lớp:', level=3)

model_table = doc.add_table(rows=1, cols=4)
model_table.style = 'Light List Accent 1'
add_table_row(model_table, ['Layer', 'Output Shape', 'Tham số', 'Chức năng'], header=True)
add_table_row(model_table, ['Input', '(32, 13, 1)', '0', 'MFCC features'])
add_table_row(model_table, ['Conv2D (8, 3×3, ReLU)', '(32, 13, 8)', '80', 'Phát hiện patterns cục bộ'])
add_table_row(model_table, ['MaxPooling2D (2×2)', '(16, 6, 8)', '0', 'Giảm chiều, giữ đặc trưng'])
add_table_row(model_table, ['Conv2D (16, 3×3, ReLU)', '(16, 6, 16)', '1,168', 'Học patterns phức tạp hơn'])
add_table_row(model_table, ['MaxPooling2D (2×2)', '(8, 3, 16)', '0', 'Giảm chiều lần 2'])
add_table_row(model_table, ['Flatten', '(384,)', '0', 'Dẹp thành vector 1D'])
add_table_row(model_table, ['Dense (32, ReLU)', '(32,)', '12,320', 'Feature vector'])
add_table_row(model_table, ['Dense (1, Sigmoid)', '(1,)', '33', 'Xác suất [0, 1]'])
add_table_row(model_table, ['Tổng', '', '13,601', '~12KB INT8'])

doc.add_paragraph()
doc.add_paragraph(
    'Tại sao chọn CNN thay vì RNN/LSTM? Vì CNN có thể phát hiện patterns cục bộ '
    'trên spectrogram, bất biến với dịch chuyển nhỏ, và TFLite Micro hỗ trợ CNN tốt hơn. '
    'Với 32 frames, CNN đã capture đủ temporal patterns.'
)

# 3.4 Training
doc.add_heading('3.4. Huấn luyện & Đánh giá', level=2)

doc.add_heading('Siêu tham số (Hyperparameters)', level=3)
hp_table = doc.add_table(rows=1, cols=3)
hp_table.style = 'Light List Accent 1'
add_table_row(hp_table, ['Tham số', 'Giá trị', 'Giải thích'], header=True)
add_table_row(hp_table, ['Optimizer', 'Adam', 'Adaptive Moment Estimation — "set-and-forget"'])
add_table_row(hp_table, ['Loss Function', 'binary_crossentropy', 'Phù hợp binary classification'])
add_table_row(hp_table, ['Metrics', 'accuracy', 'Độ chính xác (TP+TN)/total'])
add_table_row(hp_table, ['Batch Size', '16', 'Cân bằng RAM và độ ổn định gradient'])
add_table_row(hp_table, ['Epochs', '15-50', 'Dùng EarlyStopping để tự dừng'])
add_table_row(hp_table, ['Validation Split', '0.2', 'Với dataset nhỏ, cân nhắc K-Fold'])

doc.add_paragraph()
doc.add_heading('Dropout & Early Stopping (chống Overfitting)', level=3)
add_code_block(doc, """\
from tensorflow.keras.callbacks import EarlyStopping

model = Sequential([
    Input(shape=(32, 13, 1)),
    Conv2D(8, 3, activation='relu', padding='same'),
    MaxPooling2D(2),
    Conv2D(16, 3, activation='relu', padding='same'),
    MaxPooling2D(2),
    Flatten(),
    Dropout(0.3),           # ← Ngẫu nhiên tắt 30% neurons
    Dense(32, activation='relu'),
    Dropout(0.2),           # ← Ngẫu nhiên tắt 20% neurons
    Dense(1, activation='sigmoid')
])

early_stop = EarlyStopping(
    monitor='val_loss',
    patience=5,
    restore_best_weights=True
)

model.fit(X, y, epochs=50, batch_size=16,
          validation_split=0.2,
          callbacks=[early_stop])""")

doc.add_heading('Data Augmentation (tăng dataset 4×)', level=3)
add_code_block(doc, """\
def augment_audio(audio):
    aug = [audio]  # Mẫu gốc
    # 1. White noise nhẹ
    noise = np.random.normal(0, 0.005, len(audio))
    aug.append(audio + noise)
    # 2. Time shift (±100ms)
    shift = int(np.random.uniform(-1600, 1600))
    aug.append(np.roll(audio, shift))
    # 3. Volume change (±20%)
    gain = np.random.uniform(0.8, 1.2)
    aug.append(audio * gain)
    return aug  # 4 samples thay vì 1""")

doc.add_heading('Các Metrics đánh giá', level=3)
doc.add_paragraph(
    'Đối với bài toán binary classification, các metrics quan trọng:'
)

metrics_table = doc.add_table(rows=1, cols=3)
metrics_table.style = 'Light List Accent 1'
add_table_row(metrics_table, ['Metric', 'Công thức', 'Ý nghĩa'], header=True)
add_table_row(metrics_table, ['Accuracy', '(TP+TN)/(TP+TN+FP+FN)', 'Độ chính xác tổng thể'])
add_table_row(metrics_table, ['Precision', 'TP/(TP+FP)', 'Trong số mở cửa, bao nhiêu % đúng?'])
add_table_row(metrics_table, ['Recall', 'TP/(TP+FN)', 'Chủ nhà nói, bao nhiêu % được mở?'])
add_table_row(metrics_table, ['F1-Score', '2×P×R/(P+R)', 'Cân bằng Precision & Recall'])

doc.add_paragraph()
doc.add_paragraph('Ví dụ với 40 mẫu test (20 target + 20 noise):')
add_code_block(doc, """\
TP=18, TN=17, FP=3, FN=2
Accuracy  = (18+17)/40 = 87.5%
Precision = 18/(18+3)  = 85.7%
Recall    = 18/(18+2)  = 90.0%
F1-Score  = 2×0.857×0.9/(0.857+0.9) = 87.8%""")

doc.add_paragraph()
add_note(doc, 'Precision vs Recall: Bảo mật cao → ↑ threshold (giảm FP). Tiện lợi cao → ↓ threshold (giảm FN).', 'info')

# 3.5 INT8 Quantization
doc.add_heading('3.5. Lượng tử hóa INT8', level=2)

doc.add_paragraph(
    'Sau khi huấn luyện, model FP32 (~52KB) được lượng tử hóa xuống INT8 (~12KB) '
    'để chạy trên ESP32-S3. Quá trình này giảm 4× kích thước và tăng 4× tốc độ.'
)

int8_table = doc.add_table(rows=1, cols=3)
int8_table.style = 'Light List Accent 1'
add_table_row(int8_table, ['Định dạng', 'Kích thước', 'Tốc độ inference'], header=True)
add_table_row(int8_table, ['FP32', '~52KB', 'Chậm (FPU software)'])
add_table_row(int8_table, ['INT8', '~12KB', '4× nhanh hơn (PIE + SIMD)'])

doc.add_paragraph()
add_code_block(doc, """\
# Post-training INT8 quantization
def representative_data_gen():
    for x in tf.data.Dataset.from_tensor_slices(X)
                    .batch(1).take(100):
        yield [tf.cast(x, tf.float32)]

converter = TFLiteConverter.from_keras_model(model)
converter.optimizations = [tf.lite.Optimize.DEFAULT]
converter.representative_dataset = representative_data_gen
converter.target_spec.supported_ops = [TFLITE_BUILTINS_INT8]
converter.inference_input_type = tf.int8
converter.inference_output_type = tf.int8

tflite_model = converter.convert()  # ~12KB""")

# 3.6 TFLite Micro
doc.add_heading('3.6. Triển khai TFLite Micro', level=2)

doc.add_paragraph(
    'Model INT8 được chuyển thành C array và nhúng trực tiếp vào firmware:'
)

add_code_block(doc, """\
// model_data.h
extern const unsigned char g_voice_auth_model_data[];
extern const int g_voice_auth_model_data_len;

// model_data.cc (auto-generated by train_quantize_export.py)
const unsigned char g_voice_auth_model_data[] = {
  0x20, 0x00, 0x00, 0x00, 0x54, 0x46, 0x4c, 0x33, ...
};
const int g_voice_auth_model_data_len = 12000;""")

doc.add_paragraph('Khởi tạo TFLite Micro trên ESP32-S3:')
add_code_block(doc, """\
bool voice_auth_init() {
    // 1. Load model từ Flash
    model = tflite::GetModel(g_voice_auth_model_data);

    // 2. Đăng ký operators
    static MicroMutableOpResolver<8> resolver;
    resolver.AddConv2D();  resolver.AddMaxPool2D();
    resolver.AddReshape(); resolver.AddFullyConnected();
    resolver.AddLogistic(); resolver.AddShape();
    resolver.AddStridedSlice(); resolver.AddPack();

    // 3. Cấp phát tensor arena (40KB PSRAM)
    tensor_arena = (uint8_t*)heap_caps_malloc(40*1024, MALLOC_CAP_SPIRAM);

    // 4. Khởi tạo interpreter
    static MicroInterpreter interpreter(model, resolver, tensor_arena, 40*1024);
    interpreter.AllocateTensors();
    return true;
}""")

doc.add_paragraph('Pipeline inference hoàn chỉnh:')
add_code_block(doc, """\
bool voice_auth_verify(const int16_t* audio, int num_samples, float* score_out) {
    // 1. VAD: Kiểm tra có tiếng nói không
    float rms = sqrtf(sum(audio^2) / num_samples);
    if (rms < 30.0f) { *score_out = -1.0f; return false; }

    // 2. Trích MFCC (đồng bộ với Python)
    compute_mfcc_like_features(audio, num_samples, input_tensor, max_val);

    // 3. TFLite Inference
    interpreter->Invoke();

    // 4. Dequantize output (INT8 → Float)
    float prediction = (output->data.int8[0] - output->params.zero_point)
                       * output->params.scale;

    *score_out = prediction;
    return (prediction >= AUTH_THRESHOLD);  // 0.7
}""")

doc.add_page_break()

# ============================================================
# 4. FACE RECOGNITION
# ============================================================
doc.add_heading('4. Face Recognition Pipeline', level=1)

doc.add_paragraph(
    'Khác với Voice model được tự huấn luyện, Face Recognition sử dụng các model '
    'pre-trained từ thư viện ESP-DL của Espressif. Các model này đã được train sẵn '
    'và tối ưu cho ESP32-S3.'
)

doc.add_heading('4.1. Pre-trained Models (ESP-DL)', level=2)

espdl_table = doc.add_table(rows=1, cols=3)
espdl_table.style = 'Light List Accent 1'
add_table_row(espdl_table, ['Model', 'Định dạng', 'Chức năng'], header=True)
add_table_row(espdl_table, ['human_face_detect_msr_s8_v1', '.espdl', 'Multi-scale face detection (stage 1)'])
add_table_row(espdl_table, ['human_face_detect_mnp_s8_v1', '.espdl', 'Face refinement (stage 2)'])
add_table_row(espdl_table, ['MFN_S8_V1', '.espdl', 'MobileFaceNet — feature extraction'])

doc.add_paragraph()

doc.add_heading('4.2. Face Detection & Recognition', level=2)
add_code_block(doc, """\
// Khởi tạo
face_detector = new human_face_detect::MSRMNP(
    "human_face_detect_msr_s8_v1.espdl",
    0.3f, 0.3f,
    "human_face_detect_mnp_s8_v1.espdl",
    0.3f, 0.3f
);

face_recognizer = new HumanFaceRecognizer(
    "/spiflash/face.db",
    HumanFaceFeat::MFN_S8_V1,
    false  // Load database immediately
);""")

doc.add_paragraph('Pipeline:')
add_code_block(doc, """\
Camera JPEG → Decode RGB888 → Face Detection (MSRMNP)
  → Score ≥ 0.3? → Feature Extraction (MobileFaceNet)
  → Cosine Similarity ≥ 0.5? → Matched!
  Database: /spiflash/face.db (feature vectors)
  Metadata: /spiflash/face_meta.dat (tên + ID)""")

doc.add_heading('4.3. Enrollment & Database', level=2)
doc.add_paragraph(
    'Khi enroll khuôn mặt mới, feature vector (128D) được lưu vào database trên '
    'SPIFFS. Có thể enroll tối đa 10 khuôn mặt.'
)
add_code_block(doc, """\
// Enroll face mới
int face_recognition_enroll(camera_fb_t *fb, const char *name) {
    // Decode JPEG → detect face → enroll
    face_recognizer->enroll(img, detect_results);
    // Lưu metadata
    strncpy(face_database[id].name, name, MAX_NAME_LENGTH);
    save_face_metadata();
    return id;
}

// Các thresholds có thể điều chỉnh runtime:
face_recognition_set_detect_threshold(0.3);
face_recognition_set_similarity_threshold(0.5);
// Lưu trong NVS → persist khi reset!""")

doc.add_page_break()

# ============================================================
# 5. DUAL-AUTH STATE MACHINE
# ============================================================
doc.add_heading('5. Dual-Auth State Machine', level=1)

doc.add_paragraph(
    'Hệ thống sử dụng state machine với 4 trạng thái, được điều phối bởi '
    'FreeRTOS Event Groups và binary semaphore.'
)

states_table = doc.add_table(rows=1, cols=4)
states_table.style = 'Light List Accent 1'
add_table_row(states_table, ['Trạng thái', 'Mô tả', 'LED', 'Thời gian'], header=True)
add_table_row(states_table, ['STANDBY', 'Chờ Voice ID, Face OFF', 'Tắt', 'Vô thời hạn'])
add_table_row(states_table, ['VOICE_TRIGGERED', 'Voice OK, Face ON', 'Xanh lá (solid)', '~10s (chờ Face)'])
add_table_row(states_table, ['FACE_AUTH', 'Chờ Face ID result', 'Xanh lá', '≤10s (timeout)'])
add_table_row(states_table, ['DOOR_OPEN', 'Face OK, mở cửa', 'Nhấp nháy xanh', '5s'])

doc.add_paragraph()

doc.add_paragraph('Luồng hoạt động chi tiết:')
add_code_block(doc, """\
STANDBY
  → Voice matched! → VOICE_TRIGGERED (LED GREEN, Face ON)
    → Face matched! → DOOR_OPEN (servo mở, LED blink)
      → 5s → Close door → STANDBY
    → Face timeout (10s) → STANDBY
  → Voice FAIL → LED blink RED 3 lần → STANDBY""")

add_code_block(doc, """\
// Event-driven với FreeRTOS Event Groups
EventGroupHandle_t g_auth_events;
#define VOICE_AUTH_OK_EVENT   (1 << 0)
#define VOICE_AUTH_FAIL_EVENT (1 << 1)
#define FACE_AUTH_OK_EVENT    (1 << 2)
#define FACE_AUTH_FAIL_EVENT  (1 << 3)

// Trong state_machine_task():
bits = xEventGroupWaitBits(
    g_auth_events,
    VOICE_AUTH_OK_EVENT | VOICE_AUTH_FAIL_EVENT |
    FACE_AUTH_OK_EVENT | FACE_AUTH_FAIL_EVENT,
    pdTRUE, pdFALSE, portMAX_DELAY);

if (bits & VOICE_AUTH_OK_EVENT) {
    // Kích hoạt Face Recognition
    xSemaphoreGive(g_face_trigger_sem);

    // Chờ Face kết quả với timeout 10s
    bits = xEventGroupWaitBits(g_auth_events,
        FACE_AUTH_OK_EVENT | FACE_AUTH_FAIL_EVENT,
        pdTRUE, pdFALSE, pdMS_TO_TICKS(10000));

    if (bits & FACE_AUTH_OK_EVENT) {
        servo_open_door();
        vTaskDelay(pdMS_TO_TICKS(5000));
        servo_close_door();
    }
}""")

doc.add_page_break()

# ============================================================
# 6. TỐI ƯU & TROUBLESHOOTING
# ============================================================
doc.add_heading('6. Tối ưu & Troubleshooting', level=1)

doc.add_heading('6.1. Vấn đề thường gặp', level=2)

qa_table = doc.add_table(rows=1, cols=2)
qa_table.style = 'Light List Accent 1'
add_table_row(qa_table, ['Vấn đề', 'Giải pháp'], header=True)
add_table_row(qa_table, ['Train accuracy cao nhưng ESP32 inference tệ', 'MFCC không đồng bộ Python-C++. Debug: so sánh MFCC từng frame.'])
add_table_row(qa_table, ['ESP32 crash khi inference', 'Tensor arena không đủ. Tăng 40KB → 48KB.'])
add_table_row(qa_table, ['False positive (mở cửa nhầm) nhiều', 'Tăng AUTH_THRESHOLD từ 0.7 → 0.85 hoặc thêm noise samples.'])
add_table_row(qa_table, ['False negative (không nhận chủ nhà) nhiều', 'Giảm threshold 0.7 → 0.5. Thu âm lại trong nhiều điều kiện.'])

doc.add_paragraph()

doc.add_heading('6.2. Tối ưu hiệu năng', level=2)

optimizations = [
    ('PSRAM cho buffer lớn:', 'Cấp phát audio buffer (32KB) và tensor arena (40KB) trong PSRAM thay vì SRAM nội.'),
    ('Reusable buffer:', 'Tránh malloc/free mỗi frame → chống phân mảnh heap.'),
    ('Face task gated:', 'Face recognition chỉ chạy khi state machine kích hoạt, tiết kiệm năng lượng.'),
    ('Binary semaphore:', 'Face task sleep khi không dùng đến, wake ngay khi cần.'),
    ('Mutex ngắn:', 'Camera mutex giữ <5ms — chỉ copy frame, HTTP gửi sau khi release.'),
]
for name, desc in optimizations:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ' ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('6.3. Debug tools', level=2)
add_code_block(doc, """\
# Phân tích TFLite model
python tools/analyze_tflite_model.py voice_auth.tflite

# Liệt kê operators (so sánh với OpResolver)
python tools/dump_tflite_ops.py voice_auth.tflite""")

doc.add_page_break()

# ============================================================
# 7. KẾT LUẬN
# ============================================================
doc.add_heading('7. Kết Luận', level=1)

doc.add_paragraph(
    'Dự án đã xây dựng thành công hệ thống xác thực đa nhân tố trên ESP32-S3, '
    'kết hợp giữa nhận diện khuôn mặt và xác thực giọng nói. Các điểm chính:'
)

conclusions = [
    'Model Voice CNN chỉ 13.6K tham số, ~12KB sau INT8 quantization — hoàn toàn phù hợp cho vi điều khiển',
    'Đồng bộ MFCC extraction giữa Python (training) và C++ (inference) là yếu tố quyết định thành công',
    'Face Recognition sử dụng pre-trained models từ ESP-DL, không cần tự train',
    'Dual-auth (Voice + Face) tăng bảo mật đáng kể so với single-factor',
    'Toàn bộ hệ thống chạy offline, không phụ thuộc cloud, bảo vệ quyền riêng tư',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()

add_note(doc, '"Trong Machine Learning cho Embedded, dữ liệu sạch và feature engineering tốt quan trọng hơn model phức tạp. Bắt đầu từ đơn giản, đo lường, rồi mới tối ưu!"', 'success')

doc.add_page_break()

# ============================================================
# PHỤ LỤC
# ============================================================
doc.add_heading('Phụ Lục: Mã nguồn tham khảo', level=1)

doc.add_paragraph('1. Thu thập dữ liệu:', style='List Bullet')
add_code_block(doc, """\
python main/collect_data.py          # Thu target (giọng bạn)
# Đổi CLASS_NAME = "noise" rồi chạy lại
python main/collect_data.py          # Thu tiếng ồn
# Import ESC-50 dataset
python main/select_esc50.py path/to/esc50 ./esc50_selected --copy""")

doc.add_paragraph('2. Huấn luyện model:', style='List Bullet')
add_code_block(doc, """\
python main/train_quantize_export.py
# Output:
#   voice_auth.tflite  — INT8 model (~12KB)
#   model_data.cc      — C array for firmware
#   model_data.h       — Header file""")

doc.add_paragraph('3. Phân tích model:', style='List Bullet')
add_code_block(doc, """\
python tools/analyze_tflite_model.py voice_auth.tflite
python tools/dump_tflite_ops.py voice_auth.tflite""")

doc.add_paragraph('4. Cấu trúc thư mục dự án:', style='List Bullet')
add_code_block(doc, """\
esp32-s3-wifi-cam-face-recognition/
├── main/
│   ├── train_quantize_export.py   ← TRAIN MODEL (Python)
│   ├── collect_data.py            ← THU ÂM (Python)
│   ├── select_esc50.py            ← IMPORT DATA (Python)
│   ├── training_config.json       ← Lưu tham số train
│   ├── voice_auth.cpp/.h          ← Inference engine (C++)
│   ├── audio_capture.cpp/.h       ← I2S microphone driver
│   ├── face_recognition.cpp/.h    ← ESP-DL face pipeline
│   ├── main.cpp                   ← State machine + tasks
│   └── model_data.cc/.h           ← TFLite model array
├── dataset/                       ← Training data
│   ├── target/                    ← Giọng chủ nhà
│   └── noise/                     ← Tiếng ồn
├── tools/                         ← Debug tools
│   ├── analyze_tflite_model.py
│   └── dump_tflite_ops.py
├── web_assets/                    ← Web UI files
├── slides/                        ← Presentation files
└── docs/                          ← Tài liệu""")

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'Edge_AI_Model_Building_Guide.docx')
doc.save(output_path)
print(f'Da tao file: {output_path}')
print(f'Kich thuoc: {os.path.getsize(output_path) / 1024:.1f} KB')
